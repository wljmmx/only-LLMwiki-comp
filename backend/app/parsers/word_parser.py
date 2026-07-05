"""Word 解析器（基于 python-docx）

提取标题、段落、表格，保留样式层级。
"""
from __future__ import annotations

import hashlib

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.parsers.base import DocumentParser, ElementType, ParsedDocument, ParsedElement


class WordParser:
    format = "word"

    def parse(self, path: str, doc_id: str) -> ParsedDocument:
        doc = DocxDocument(path)

        with open(path, "rb") as f:
            checksum = hashlib.sha256(f.read()).hexdigest()

        title = doc.paragraphs[0].text.strip() if doc.paragraphs else None
        elements = self._parse_docx(doc)

        return ParsedDocument(
            doc_id=doc_id, source_path=path, format="word",
            checksum=checksum, title=title, elements=elements,
        )

    def _parse_docx(self, doc: DocxDocument) -> list[ParsedElement]:
        elements: list[ParsedElement] = []
        current_section: str | None = None

        # 遍历文档 body 中的段落和表格
        from docx.oxml.ns import qn
        body = doc.element.body

        for child in body:
            if child.tag == qn("w:p"):
                para = self._find_paragraph(doc, child)
                if para is None:
                    continue
                text = para.text.strip()
                if not text:
                    continue
                style = para.style.name if para.style else ""

                if "Heading" in style or "heading" in style.lower() or "标题" in style:
                    level = 1
                    for s in style.split():
                        if s.isdigit():
                            level = min(int(s), 6)
                    elements.append(ParsedElement(
                        type=ElementType.HEADING, content=text,
                        metadata={"level": level, "style": style},
                    ))
                    if level <= 2:
                        current_section = text
                else:
                    elements.append(ParsedElement(
                        type=ElementType.PARAGRAPH, content=text,
                        section=current_section,
                        metadata={"style": style},
                    ))

            elif child.tag == qn("w:tbl"):
                table = self._find_table(doc, child)
                if table is not None:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(" | ".join(cells))
                    if rows:
                        elements.append(ParsedElement(
                            type=ElementType.TABLE, content="\n".join(rows),
                            section=current_section,
                        ))

        return elements

    def _find_paragraph(self, doc, element):
        """根据 XML 元素找到对应的 Paragraph 对象"""
        for para in doc.paragraphs:
            if para._element is element:
                return para
        return None

    def _find_table(self, doc, element):
        for table in doc.tables:
            if table._element is element:
                return table
        return None